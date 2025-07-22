#!/usr/bin/env python3
"""
Metadata Extractor Service for the AIMA Media Lifecycle Management Service.

This module provides comprehensive metadata extraction capabilities for various
media file types including images, videos, audio files, and documents.
"""

import logging
import os
import mimetypes
from typing import Any, Dict, List, Optional, Tuple, Union
from pathlib import Path
from datetime import datetime
from enum import Enum
from dataclasses import dataclass, asdict
import hashlib
import json
from uuid import UUID, uuid4

# Media processing libraries
try:
    from PIL import Image, ExifTags
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("PIL/Pillow not available - image metadata extraction disabled")

try:
    import cv2
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False
    logger.warning("OpenCV not available - advanced video analysis disabled")

try:
    from mutagen import File as MutagenFile
    from mutagen.id3 import ID3NoHeaderError
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False
    logger.warning("Mutagen not available - audio metadata extraction disabled")

try:
    import ffmpeg
    FFMPEG_AVAILABLE = True
except ImportError:
    FFMPEG_AVAILABLE = False
    logger.warning("ffmpeg-python not available - video metadata extraction limited")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not available - PDF metadata extraction disabled")

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX metadata extraction disabled")

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, or_, func

from ..core.database import MediaFile, MediaMetadata
from ..core.redis_client import CacheManager
from ..middleware.error_handling import MediaProcessingError
from .audit_service import AuditService, AuditEventType, AuditSeverity


logger = logging.getLogger(__name__)


class MetadataType(str, Enum):
    """Types of metadata that can be extracted."""
    BASIC = "basic"
    EXIF = "exif"
    TECHNICAL = "technical"
    DESCRIPTIVE = "descriptive"
    ADMINISTRATIVE = "administrative"
    STRUCTURAL = "structural"
    PRESERVATION = "preservation"
    RIGHTS = "rights"


class MediaType(str, Enum):
    """Supported media types for metadata extraction."""
    IMAGE = "image"
    VIDEO = "video"
    AUDIO = "audio"
    DOCUMENT = "document"
    ARCHIVE = "archive"
    OTHER = "other"


@dataclass
class BasicMetadata:
    """Basic file metadata."""
    filename: str
    file_size: int
    file_type: str
    mime_type: str
    created_at: Optional[datetime] = None
    modified_at: Optional[datetime] = None
    checksum_md5: Optional[str] = None
    checksum_sha256: Optional[str] = None


@dataclass
class ImageMetadata:
    """Image-specific metadata."""
    width: Optional[int] = None
    height: Optional[int] = None
    color_mode: Optional[str] = None
    bit_depth: Optional[int] = None
    compression: Optional[str] = None
    orientation: Optional[int] = None
    dpi: Optional[Tuple[float, float]] = None
    has_transparency: Optional[bool] = None
    color_profile: Optional[str] = None
    
    # EXIF data
    camera_make: Optional[str] = None
    camera_model: Optional[str] = None
    lens_model: Optional[str] = None
    focal_length: Optional[float] = None
    aperture: Optional[float] = None
    shutter_speed: Optional[str] = None
    iso: Optional[int] = None
    flash: Optional[str] = None
    white_balance: Optional[str] = None
    exposure_mode: Optional[str] = None
    metering_mode: Optional[str] = None
    
    # GPS data
    gps_latitude: Optional[float] = None
    gps_longitude: Optional[float] = None
    gps_altitude: Optional[float] = None
    gps_timestamp: Optional[datetime] = None
    
    # Date/time
    date_taken: Optional[datetime] = None
    date_digitized: Optional[datetime] = None
    date_modified: Optional[datetime] = None


@dataclass
class VideoMetadata:
    """Video-specific metadata."""
    width: Optional[int] = None
    height: Optional[int] = None
    duration: Optional[float] = None
    frame_rate: Optional[float] = None
    bit_rate: Optional[int] = None
    codec: Optional[str] = None
    container_format: Optional[str] = None
    aspect_ratio: Optional[str] = None
    color_space: Optional[str] = None
    
    # Audio track info
    audio_codec: Optional[str] = None
    audio_bit_rate: Optional[int] = None
    audio_sample_rate: Optional[int] = None
    audio_channels: Optional[int] = None
    
    # Metadata
    title: Optional[str] = None
    description: Optional[str] = None
    creation_time: Optional[datetime] = None
    
    # Technical details
    total_frames: Optional[int] = None
    keyframe_interval: Optional[int] = None
    has_subtitles: Optional[bool] = None
    subtitle_languages: Optional[List[str]] = None


@dataclass
class AudioMetadata:
    """Audio-specific metadata."""
    duration: Optional[float] = None
    bit_rate: Optional[int] = None
    sample_rate: Optional[int] = None
    channels: Optional[int] = None
    codec: Optional[str] = None
    
    # ID3 tags
    title: Optional[str] = None
    artist: Optional[str] = None
    album: Optional[str] = None
    album_artist: Optional[str] = None
    genre: Optional[str] = None
    year: Optional[int] = None
    track_number: Optional[int] = None
    total_tracks: Optional[int] = None
    disc_number: Optional[int] = None
    total_discs: Optional[int] = None
    composer: Optional[str] = None
    comment: Optional[str] = None
    
    # Technical details
    encoding: Optional[str] = None
    variable_bitrate: Optional[bool] = None
    lossless: Optional[bool] = None


@dataclass
class DocumentMetadata:
    """Document-specific metadata."""
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    producer: Optional[str] = None
    creation_date: Optional[datetime] = None
    modification_date: Optional[datetime] = None
    keywords: Optional[List[str]] = None
    
    # Content details
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    language: Optional[str] = None
    
    # Security
    encrypted: Optional[bool] = None
    permissions: Optional[Dict[str, bool]] = None


@dataclass
class ExtractedMetadata:
    """Complete extracted metadata for a media file."""
    file_id: UUID
    extraction_id: UUID
    basic: BasicMetadata
    image: Optional[ImageMetadata] = None
    video: Optional[VideoMetadata] = None
    audio: Optional[AudioMetadata] = None
    document: Optional[DocumentMetadata] = None
    
    # Additional metadata
    custom_metadata: Optional[Dict[str, Any]] = None
    extraction_timestamp: Optional[datetime] = None
    extraction_duration: Optional[float] = None
    extraction_errors: Optional[List[str]] = None
    
    def __post_init__(self):
        if self.extraction_timestamp is None:
            self.extraction_timestamp = datetime.utcnow()
        if self.extraction_id is None:
            self.extraction_id = uuid4()


class MetadataExtractor:
    """Service for extracting metadata from media files."""
    
    def __init__(
        self,
        cache_manager: CacheManager,
        audit_service: Optional[AuditService] = None
    ):
        self.cache = cache_manager
        self.audit_service = audit_service
        
        # Configuration
        self.cache_ttl = 3600  # 1 hour
        self.max_file_size = 500 * 1024 * 1024  # 500MB
        self.supported_image_formats = {
            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif',
            '.webp', '.ico', '.svg', '.raw', '.cr2', '.nef', '.arw'
        }
        self.supported_video_formats = {
            '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm',
            '.m4v', '.3gp', '.ogv', '.ts', '.mts', '.m2ts'
        }
        self.supported_audio_formats = {
            '.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a',
            '.opus', '.ape', '.ac3', '.dts'
        }
        self.supported_document_formats = {
            '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
            '.txt', '.rtf', '.odt', '.ods', '.odp'
        }
        
        # Initialize extractors
        self._initialize_extractors()
    
    def _initialize_extractors(self):
        """Initialize available metadata extractors."""
        self.available_extractors = {
            'basic': True,
            'image_pil': PIL_AVAILABLE,
            'video_opencv': OPENCV_AVAILABLE,
            'video_ffmpeg': FFMPEG_AVAILABLE,
            'audio_mutagen': MUTAGEN_AVAILABLE,
            'document_pypdf2': PYPDF2_AVAILABLE,
            'document_docx': PYTHON_DOCX_AVAILABLE
        }
        
        logger.info(f"Metadata extractors initialized: {self.available_extractors}")
    
    async def extract_metadata(
        self,
        db: AsyncSession,
        file_path: str,
        file_id: Optional[UUID] = None,
        force_refresh: bool = False
    ) -> ExtractedMetadata:
        """
        Extract comprehensive metadata from a media file.
        
        Args:
            db: Database session
            file_path: Path to the media file
            file_id: Optional file ID for caching
            force_refresh: Force re-extraction even if cached
        
        Returns:
            Extracted metadata
        """
        start_time = datetime.utcnow()
        extraction_errors = []
        
        try:
            # Check cache first
            if file_id and not force_refresh:
                cached_metadata = await self._get_cached_metadata(file_id)
                if cached_metadata:
                    return cached_metadata
            
            # Validate file
            if not os.path.exists(file_path):
                raise MediaProcessingError(f"File not found: {file_path}")
            
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                raise MediaProcessingError(f"File too large: {file_size} bytes")
            
            # Extract basic metadata
            basic_metadata = await self._extract_basic_metadata(file_path)
            
            # Determine media type
            media_type = self._determine_media_type(file_path, basic_metadata.mime_type)
            
            # Extract type-specific metadata
            image_metadata = None
            video_metadata = None
            audio_metadata = None
            document_metadata = None
            
            try:
                if media_type == MediaType.IMAGE:
                    image_metadata = await self._extract_image_metadata(file_path)
                elif media_type == MediaType.VIDEO:
                    video_metadata = await self._extract_video_metadata(file_path)
                elif media_type == MediaType.AUDIO:
                    audio_metadata = await self._extract_audio_metadata(file_path)
                elif media_type == MediaType.DOCUMENT:
                    document_metadata = await self._extract_document_metadata(file_path)
            except Exception as e:
                extraction_errors.append(f"Type-specific extraction failed: {str(e)}")
                logger.warning(f"Failed to extract {media_type.value} metadata: {e}")
            
            # Create extracted metadata object
            extraction_duration = (datetime.utcnow() - start_time).total_seconds()
            
            extracted_metadata = ExtractedMetadata(
                file_id=file_id or uuid4(),
                extraction_id=uuid4(),
                basic=basic_metadata,
                image=image_metadata,
                video=video_metadata,
                audio=audio_metadata,
                document=document_metadata,
                extraction_timestamp=start_time,
                extraction_duration=extraction_duration,
                extraction_errors=extraction_errors if extraction_errors else None
            )
            
            # Cache the metadata
            if file_id:
                await self._cache_metadata(file_id, extracted_metadata)
            
            # Log extraction
            if self.audit_service:
                from .audit_service import AuditContext, AuditEvent
                
                context = AuditContext()
                audit_event = AuditEvent(
                    event_type=AuditEventType.MEDIA_PROCESSED,
                    severity=AuditSeverity.LOW,
                    status="success",
                    message=f"Metadata extracted from {os.path.basename(file_path)}",
                    context=context,
                    details={
                        "file_path": file_path,
                        "media_type": media_type.value,
                        "extraction_duration": extraction_duration,
                        "errors": extraction_errors
                    }
                )
                
                await self.audit_service.log_event(db, audit_event)
            
            return extracted_metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction failed for {file_path}: {e}")
            
            # Log extraction failure
            if self.audit_service:
                from .audit_service import AuditContext, AuditEvent
                
                context = AuditContext()
                audit_event = AuditEvent(
                    event_type=AuditEventType.MEDIA_PROCESSING_FAILED,
                    severity=AuditSeverity.HIGH,
                    status="failure",
                    message=f"Metadata extraction failed for {os.path.basename(file_path)}",
                    context=context,
                    details={"file_path": file_path, "error": str(e)}
                )
                
                await self.audit_service.log_event(db, audit_event)
            
            raise MediaProcessingError(f"Metadata extraction failed: {str(e)}")
    
    async def _extract_basic_metadata(self, file_path: str) -> BasicMetadata:
        """Extract basic file metadata."""
        try:
            file_stat = os.stat(file_path)
            filename = os.path.basename(file_path)
            
            # Determine MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = 'application/octet-stream'
            
            # Calculate checksums
            md5_hash = hashlib.md5()
            sha256_hash = hashlib.sha256()
            
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    md5_hash.update(chunk)
                    sha256_hash.update(chunk)
            
            return BasicMetadata(
                filename=filename,
                file_size=file_stat.st_size,
                file_type=Path(file_path).suffix.lower(),
                mime_type=mime_type,
                created_at=datetime.fromtimestamp(file_stat.st_ctime),
                modified_at=datetime.fromtimestamp(file_stat.st_mtime),
                checksum_md5=md5_hash.hexdigest(),
                checksum_sha256=sha256_hash.hexdigest()
            )
            
        except Exception as e:
            logger.error(f"Basic metadata extraction failed: {e}")
            raise MediaProcessingError(f"Failed to extract basic metadata: {str(e)}")
    
    async def _extract_image_metadata(self, file_path: str) -> Optional[ImageMetadata]:
        """Extract image-specific metadata."""
        if not PIL_AVAILABLE:
            return None
        
        try:
            with Image.open(file_path) as img:
                metadata = ImageMetadata(
                    width=img.width,
                    height=img.height,
                    color_mode=img.mode,
                    has_transparency='transparency' in img.info
                )
                
                # Get DPI information
                if hasattr(img, 'info') and 'dpi' in img.info:
                    metadata.dpi = img.info['dpi']
                
                # Extract EXIF data
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    metadata = self._parse_exif_data(metadata, exif_data)
                
                # Get color profile
                if hasattr(img, 'info') and 'icc_profile' in img.info:
                    metadata.color_profile = 'ICC Profile present'
                
                return metadata
                
        except Exception as e:
            logger.warning(f"Image metadata extraction failed: {e}")
            return None
    
    def _parse_exif_data(self, metadata: ImageMetadata, exif_data: Dict) -> ImageMetadata:
        """Parse EXIF data and update image metadata."""
        try:
            for tag_id, value in exif_data.items():
                tag = TAGS.get(tag_id, tag_id)
                
                if tag == 'Make':
                    metadata.camera_make = str(value)
                elif tag == 'Model':
                    metadata.camera_model = str(value)
                elif tag == 'LensModel':
                    metadata.lens_model = str(value)
                elif tag == 'FocalLength':
                    if isinstance(value, tuple) and len(value) == 2:
                        metadata.focal_length = float(value[0]) / float(value[1])
                elif tag == 'FNumber':
                    if isinstance(value, tuple) and len(value) == 2:
                        metadata.aperture = float(value[0]) / float(value[1])
                elif tag == 'ExposureTime':
                    if isinstance(value, tuple) and len(value) == 2:
                        metadata.shutter_speed = f"1/{int(value[1]/value[0])}"
                elif tag == 'ISOSpeedRatings':
                    metadata.iso = int(value)
                elif tag == 'Flash':
                    metadata.flash = self._decode_flash_value(value)
                elif tag == 'WhiteBalance':
                    metadata.white_balance = 'Auto' if value == 0 else 'Manual'
                elif tag == 'ExposureMode':
                    exposure_modes = {0: 'Auto', 1: 'Manual', 2: 'Auto bracket'}
                    metadata.exposure_mode = exposure_modes.get(value, 'Unknown')
                elif tag == 'MeteringMode':
                    metering_modes = {
                        1: 'Average', 2: 'Center-weighted', 3: 'Spot',
                        4: 'Multi-spot', 5: 'Pattern', 6: 'Partial'
                    }
                    metadata.metering_mode = metering_modes.get(value, 'Unknown')
                elif tag == 'Orientation':
                    metadata.orientation = int(value)
                elif tag == 'DateTime':
                    try:
                        metadata.date_taken = datetime.strptime(str(value), '%Y:%m:%d %H:%M:%S')
                    except ValueError:
                        pass
                elif tag == 'DateTimeDigitized':
                    try:
                        metadata.date_digitized = datetime.strptime(str(value), '%Y:%m:%d %H:%M:%S')
                    except ValueError:
                        pass
                elif tag == 'GPSInfo':
                    metadata = self._parse_gps_data(metadata, value)
            
            return metadata
            
        except Exception as e:
            logger.warning(f"EXIF parsing failed: {e}")
            return metadata
    
    def _parse_gps_data(self, metadata: ImageMetadata, gps_info: Dict) -> ImageMetadata:
        """Parse GPS data from EXIF."""
        try:
            if not gps_info:
                return metadata
            
            # Convert GPS coordinates
            if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
                lat = self._convert_gps_coordinate(gps_info['GPSLatitude'])
                if gps_info['GPSLatitudeRef'] == 'S':
                    lat = -lat
                metadata.gps_latitude = lat
            
            if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
                lon = self._convert_gps_coordinate(gps_info['GPSLongitude'])
                if gps_info['GPSLongitudeRef'] == 'W':
                    lon = -lon
                metadata.gps_longitude = lon
            
            if 'GPSAltitude' in gps_info:
                altitude = gps_info['GPSAltitude']
                if isinstance(altitude, tuple) and len(altitude) == 2:
                    metadata.gps_altitude = float(altitude[0]) / float(altitude[1])
            
            # GPS timestamp
            if 'GPSTimeStamp' in gps_info and 'GPSDateStamp' in gps_info:
                try:
                    time_stamp = gps_info['GPSTimeStamp']
                    date_stamp = gps_info['GPSDateStamp']
                    
                    if isinstance(time_stamp, tuple) and len(time_stamp) == 3:
                        hours = float(time_stamp[0][0]) / float(time_stamp[0][1]) if isinstance(time_stamp[0], tuple) else time_stamp[0]
                        minutes = float(time_stamp[1][0]) / float(time_stamp[1][1]) if isinstance(time_stamp[1], tuple) else time_stamp[1]
                        seconds = float(time_stamp[2][0]) / float(time_stamp[2][1]) if isinstance(time_stamp[2], tuple) else time_stamp[2]
                        
                        gps_datetime_str = f"{date_stamp} {int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}"
                        metadata.gps_timestamp = datetime.strptime(gps_datetime_str, '%Y:%m:%d %H:%M:%S')
                except Exception as e:
                    logger.debug(f"GPS timestamp parsing failed: {e}")
            
            return metadata
            
        except Exception as e:
            logger.warning(f"GPS data parsing failed: {e}")
            return metadata
    
    def _convert_gps_coordinate(self, coord_tuple: Tuple) -> float:
        """Convert GPS coordinate from tuple to decimal degrees."""
        if len(coord_tuple) != 3:
            return 0.0
        
        degrees = float(coord_tuple[0][0]) / float(coord_tuple[0][1]) if isinstance(coord_tuple[0], tuple) else coord_tuple[0]
        minutes = float(coord_tuple[1][0]) / float(coord_tuple[1][1]) if isinstance(coord_tuple[1], tuple) else coord_tuple[1]
        seconds = float(coord_tuple[2][0]) / float(coord_tuple[2][1]) if isinstance(coord_tuple[2], tuple) else coord_tuple[2]
        
        return degrees + (minutes / 60.0) + (seconds / 3600.0)
    
    def _decode_flash_value(self, flash_value: int) -> str:
        """Decode flash EXIF value."""
        flash_modes = {
            0: 'No Flash',
            1: 'Flash',
            5: 'Flash, No Return',
            7: 'Flash, Return',
            9: 'Flash, Compulsory',
            13: 'Flash, Compulsory, No Return',
            15: 'Flash, Compulsory, Return',
            16: 'No Flash, Compulsory',
            24: 'No Flash, Auto',
            25: 'Flash, Auto',
            29: 'Flash, Auto, No Return',
            31: 'Flash, Auto, Return',
            32: 'No Flash Available'
        }
        return flash_modes.get(flash_value, f'Unknown ({flash_value})')
    
    async def _extract_video_metadata(self, file_path: str) -> Optional[VideoMetadata]:
        """Extract video-specific metadata."""
        metadata = VideoMetadata()
        
        # Try ffmpeg first (more comprehensive)
        if FFMPEG_AVAILABLE:
            try:
                probe = ffmpeg.probe(file_path)
                
                # Get video stream info
                video_stream = next((stream for stream in probe['streams'] if stream['codec_type'] == 'video'), None)
                if video_stream:
                    metadata.width = int(video_stream.get('width', 0))
                    metadata.height = int(video_stream.get('height', 0))
                    metadata.codec = video_stream.get('codec_name')
                    metadata.bit_rate = int(video_stream.get('bit_rate', 0)) if video_stream.get('bit_rate') else None
                    
                    # Frame rate
                    if 'r_frame_rate' in video_stream:
                        frame_rate_str = video_stream['r_frame_rate']
                        if '/' in frame_rate_str:
                            num, den = frame_rate_str.split('/')
                            metadata.frame_rate = float(num) / float(den)
                    
                    # Duration
                    if 'duration' in video_stream:
                        metadata.duration = float(video_stream['duration'])
                    
                    # Aspect ratio
                    if metadata.width and metadata.height:
                        from fractions import Fraction
                        ratio = Fraction(metadata.width, metadata.height)
                        metadata.aspect_ratio = f"{ratio.numerator}:{ratio.denominator}"
                
                # Get audio stream info
                audio_stream = next((stream for stream in probe['streams'] if stream['codec_type'] == 'audio'), None)
                if audio_stream:
                    metadata.audio_codec = audio_stream.get('codec_name')
                    metadata.audio_bit_rate = int(audio_stream.get('bit_rate', 0)) if audio_stream.get('bit_rate') else None
                    metadata.audio_sample_rate = int(audio_stream.get('sample_rate', 0)) if audio_stream.get('sample_rate') else None
                    metadata.audio_channels = int(audio_stream.get('channels', 0)) if audio_stream.get('channels') else None
                
                # Get format info
                format_info = probe.get('format', {})
                metadata.container_format = format_info.get('format_name')
                if not metadata.duration and 'duration' in format_info:
                    metadata.duration = float(format_info['duration'])
                
                # Get metadata tags
                tags = format_info.get('tags', {})
                metadata.title = tags.get('title')
                metadata.description = tags.get('description') or tags.get('comment')
                
                if 'creation_time' in tags:
                    try:
                        metadata.creation_time = datetime.fromisoformat(tags['creation_time'].replace('Z', '+00:00'))
                    except ValueError:
                        pass
                
                return metadata
                
            except Exception as e:
                logger.warning(f"FFmpeg video metadata extraction failed: {e}")
        
        # Fallback to OpenCV
        if OPENCV_AVAILABLE:
            try:
                cap = cv2.VideoCapture(file_path)
                
                if cap.isOpened():
                    metadata.width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                    metadata.height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                    metadata.frame_rate = cap.get(cv2.CAP_PROP_FPS)
                    
                    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                    if frame_count > 0 and metadata.frame_rate > 0:
                        metadata.duration = frame_count / metadata.frame_rate
                        metadata.total_frames = frame_count
                
                cap.release()
                return metadata
                
            except Exception as e:
                logger.warning(f"OpenCV video metadata extraction failed: {e}")
        
        return metadata if any(getattr(metadata, field) is not None for field in metadata.__dataclass_fields__) else None
    
    async def _extract_audio_metadata(self, file_path: str) -> Optional[AudioMetadata]:
        """Extract audio-specific metadata."""
        if not MUTAGEN_AVAILABLE:
            return None
        
        try:
            audio_file = MutagenFile(file_path)
            if not audio_file:
                return None
            
            metadata = AudioMetadata()
            
            # Basic audio properties
            if hasattr(audio_file, 'info'):
                info = audio_file.info
                metadata.duration = getattr(info, 'length', None)
                metadata.bit_rate = getattr(info, 'bitrate', None)
                metadata.sample_rate = getattr(info, 'sample_rate', None)
                metadata.channels = getattr(info, 'channels', None)
                metadata.codec = getattr(info, 'codec', None)
                metadata.variable_bitrate = getattr(info, 'sketchy', None)
            
            # Extract tags
            if audio_file.tags:
                tags = audio_file.tags
                
                # Common tag mappings
                tag_mappings = {
                    'TIT2': 'title', 'TITLE': 'title', '\xa9nam': 'title',
                    'TPE1': 'artist', 'ARTIST': 'artist', '\xa9ART': 'artist',
                    'TALB': 'album', 'ALBUM': 'album', '\xa9alb': 'album',
                    'TPE2': 'album_artist', 'ALBUMARTIST': 'album_artist', 'aART': 'album_artist',
                    'TCON': 'genre', 'GENRE': 'genre', '\xa9gen': 'genre',
                    'TDRC': 'year', 'DATE': 'year', '\xa9day': 'year',
                    'TRCK': 'track_number', 'TRACKNUMBER': 'track_number', 'trkn': 'track_number',
                    'TPOS': 'disc_number', 'DISCNUMBER': 'disc_number', 'disk': 'disc_number',
                    'TCOM': 'composer', 'COMPOSER': 'composer', '\xa9wrt': 'composer',
                    'COMM': 'comment', 'COMMENT': 'comment', '\xa9cmt': 'comment'
                }
                
                for tag_key, metadata_key in tag_mappings.items():
                    if tag_key in tags:
                        value = tags[tag_key]
                        if isinstance(value, list) and value:
                            value = value[0]
                        
                        if metadata_key in ['year', 'track_number', 'disc_number']:
                            try:
                                # Extract number from string if needed
                                if isinstance(value, str):
                                    import re
                                    numbers = re.findall(r'\d+', str(value))
                                    if numbers:
                                        value = int(numbers[0])
                                else:
                                    value = int(value)
                                setattr(metadata, metadata_key, value)
                            except (ValueError, TypeError):
                                pass
                        else:
                            setattr(metadata, metadata_key, str(value))
                
                # Handle track/total tracks format (e.g., "3/12")
                if metadata.track_number and isinstance(metadata.track_number, str) and '/' in str(metadata.track_number):
                    try:
                        track_parts = str(metadata.track_number).split('/')
                        metadata.track_number = int(track_parts[0])
                        if len(track_parts) > 1:
                            metadata.total_tracks = int(track_parts[1])
                    except ValueError:
                        pass
            
            return metadata
            
        except ID3NoHeaderError:
            logger.debug(f"No ID3 header found in {file_path}")
            return None
        except Exception as e:
            logger.warning(f"Audio metadata extraction failed: {e}")
            return None
    
    async def _extract_document_metadata(self, file_path: str) -> Optional[DocumentMetadata]:
        """Extract document-specific metadata."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._extract_pdf_metadata(file_path)
        elif file_ext in ['.doc', '.docx']:
            return await self._extract_docx_metadata(file_path)
        else:
            return None
    
    async def _extract_pdf_metadata(self, file_path: str) -> Optional[DocumentMetadata]:
        """Extract PDF metadata."""
        if not PYPDF2_AVAILABLE:
            return None
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = DocumentMetadata()
                metadata.page_count = len(pdf_reader.pages)
                metadata.encrypted = pdf_reader.is_encrypted
                
                # Extract document info
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    metadata.title = info.get('/Title')
                    metadata.author = info.get('/Author')
                    metadata.subject = info.get('/Subject')
                    metadata.creator = info.get('/Creator')
                    metadata.producer = info.get('/Producer')
                    
                    # Parse dates
                    creation_date = info.get('/CreationDate')
                    if creation_date:
                        try:
                            # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
                            date_str = creation_date.replace('D:', '').split('+')[0].split('-')[0]
                            if len(date_str) >= 14:
                                metadata.creation_date = datetime.strptime(date_str[:14], '%Y%m%d%H%M%S')
                        except ValueError:
                            pass
                    
                    mod_date = info.get('/ModDate')
                    if mod_date:
                        try:
                            date_str = mod_date.replace('D:', '').split('+')[0].split('-')[0]
                            if len(date_str) >= 14:
                                metadata.modification_date = datetime.strptime(date_str[:14], '%Y%m%d%H%M%S')
                        except ValueError:
                            pass
                
                return metadata
                
        except Exception as e:
            logger.warning(f"PDF metadata extraction failed: {e}")
            return None
    
    async def _extract_docx_metadata(self, file_path: str) -> Optional[DocumentMetadata]:
        """Extract DOCX metadata."""
        if not PYTHON_DOCX_AVAILABLE:
            return None
        
        try:
            doc = DocxDocument(file_path)
            
            metadata = DocumentMetadata()
            
            # Core properties
            if doc.core_properties:
                props = doc.core_properties
                metadata.title = props.title
                metadata.author = props.author
                metadata.subject = props.subject
                metadata.creator = props.creator
                metadata.creation_date = props.created
                metadata.modification_date = props.modified
                
                if props.keywords:
                    metadata.keywords = [kw.strip() for kw in props.keywords.split(',')]
            
            # Count paragraphs and estimate word count
            paragraph_count = len(doc.paragraphs)
            word_count = 0
            character_count = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                character_count += len(text)
                word_count += len(text.split())
            
            metadata.word_count = word_count
            metadata.character_count = character_count
            
            return metadata
            
        except Exception as e:
            logger.warning(f"DOCX metadata extraction failed: {e}")
            return None
    
    def _determine_media_type(self, file_path: str, mime_type: str) -> MediaType:
        """Determine the media type of a file."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in self.supported_image_formats or mime_type.startswith('image/'):
            return MediaType.IMAGE
        elif file_ext in self.supported_video_formats or mime_type.startswith('video/'):
            return MediaType.VIDEO
        elif file_ext in self.supported_audio_formats or mime_type.startswith('audio/'):
            return MediaType.AUDIO
        elif file_ext in self.supported_document_formats or mime_type.startswith('application/'):
            return MediaType.DOCUMENT
        else:
            return MediaType.OTHER
    
    async def _get_cached_metadata(self, file_id: UUID) -> Optional[ExtractedMetadata]:
        """Get cached metadata for a file."""
        try:
            cache_key = f"metadata:{file_id}"
            cached_data = await self.cache.get(cache_key)
            
            if cached_data:
                # Deserialize the cached metadata
                return ExtractedMetadata(**cached_data)
            
            return None
            
        except Exception as e:
            logger.warning(f"Failed to get cached metadata: {e}")
            return None
    
    async def _cache_metadata(self, file_id: UUID, metadata: ExtractedMetadata):
        """Cache extracted metadata."""
        try:
            cache_key = f"metadata:{file_id}"
            # Serialize metadata to dict
            metadata_dict = asdict(metadata)
            
            # Convert datetime objects to ISO strings for JSON serialization
            def convert_datetime(obj):
                if isinstance(obj, datetime):
                    return obj.isoformat()
                elif isinstance(obj, dict):
                    return {k: convert_datetime(v) for k, v in obj.items()}
                elif isinstance(obj, list):
                    return [convert_datetime(item) for item in obj]
                return obj
            
            serializable_dict = convert_datetime(metadata_dict)
            
            await self.cache.set(cache_key, serializable_dict, ttl=self.cache_ttl)
            
        except Exception as e:
            logger.warning(f"Failed to cache metadata: {e}")
    
    async def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats by category."""
        return {
            'image': list(self.supported_image_formats),
            'video': list(self.supported_video_formats),
            'audio': list(self.supported_audio_formats),
            'document': list(self.supported_document_formats)
        }
    
    async def validate_file_support(self, file_path: str) -> Tuple[bool, str]:
        """Validate if a file is supported for metadata extraction."""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            all_supported = (
                self.supported_image_formats |
                self.supported_video_formats |
                self.supported_audio_formats |
                self.supported_document_formats
            )
            
            if file_ext in all_supported:
                return True, "File format is supported"
            else:
                return False, f"Unsupported file format: {file_ext}"
                
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    async def health_check(self) -> Dict[str, Any]:
        """Perform health check of the metadata extractor."""
        return {
            "status": "healthy",
            "available_extractors": self.available_extractors,
            "supported_formats": await self.get_supported_formats(),
            "max_file_size": self.max_file_size,
            "cache_ttl": self.cache_ttl,
            "timestamp": datetime.utcnow().isoformat()
        }